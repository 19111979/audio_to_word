import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import threading
import os
import sys
from pathlib import Path
from docx import Document
from datetime import datetime
import tempfile


class AudioToWordConverter:
    def __init__(self, root):
        self.root = root
        self.root.title("Conversor de Áudio para Word - Whisper Local")
        self.root.geometry("900x700")
        
        self.audio_path = None
        self.transcription_text = ""
        
        self.create_widgets()
    
    def create_widgets(self):
        # Frame principal
        main_frame = ttk.Frame(self.root, padding="15")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Título
        title_label = ttk.Label(main_frame, text="Conversor de Áudio para Word - Whisper Local", 
                               font=("Arial", 16, "bold"))
        title_label.grid(row=0, column=0, columnspan=3, pady=(0, 15))
        
        # Informações
        info_text = ("Este programa usa OpenAI Whisper local para transcrever áudio.\n"
                    "Não é necessário API key. Formatos suportados: MP3, WAV, M4A, FLAC, AAC, etc.")
        info_label = ttk.Label(main_frame, text=info_text, justify=tk.CENTER)
        info_label.grid(row=1, column=0, columnspan=3, pady=(0, 10))
        
        # Seleção de arquivo
        file_frame = ttk.LabelFrame(main_frame, text="Seleção de Arquivo", padding="10")
        file_frame.grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        
        self.file_label = ttk.Label(file_frame, text="Nenhum arquivo selecionado")
        self.file_label.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 10))
        
        self.select_btn = ttk.Button(file_frame, text="Selecionar Áudio", 
                                   command=self.select_audio_file)
        self.select_btn.grid(row=0, column=1)
        
        # Configurações
        settings_frame = ttk.LabelFrame(main_frame, text="Configurações", padding="10")
        settings_frame.grid(row=3, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        
        ttk.Label(settings_frame, text="Modelo Whisper:").grid(row=0, column=0, sticky=tk.W)
        self.model_var = tk.StringVar(value="base")
        model_combo = ttk.Combobox(settings_frame, textvariable=self.model_var, 
                                  values=["tiny", "base", "small", "medium", "large"],
                                  state="readonly", width=10)
        model_combo.grid(row=0, column=1, sticky=tk.W, padx=(5, 20))
        
        ttk.Label(settings_frame, text="Idioma (opcional):").grid(row=0, column=2, sticky=tk.W)
        self.language_var = tk.StringVar(value="")
        language_combo = ttk.Combobox(settings_frame, textvariable=self.language_var,
                                     values=["", "pt", "en", "es", "fr", "de", "it"],
                                     state="normal", width=5)
        language_combo.grid(row=0, column=3, sticky=tk.W, padx=5)
        
        # Botões de ação
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=4, column=0, columnspan=3, pady=15)
        
        self.transcribe_btn = ttk.Button(button_frame, text="🎤 Transcrever Áudio",
                                       command=self.start_transcription,
                                       style="Accent.TButton")
        self.transcribe_btn.grid(row=0, column=0, padx=5)
        
        self.export_btn = ttk.Button(button_frame, text="📄 Exportar para Word",
                                   command=self.export_to_word)
        self.export_btn.grid(row=0, column=1, padx=5)
        
        self.clear_btn = ttk.Button(button_frame, text="🗑️ Limpar",
                                  command=self.clear_all)
        self.clear_btn.grid(row=0, column=2, padx=5)
        
        # Barra de progresso
        self.progress = ttk.Progressbar(main_frame, mode='determinate')
        self.progress.grid(row=5, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        self.progress.grid_remove()
        
        # Área de texto
        text_frame = ttk.LabelFrame(main_frame, text="Transcrição", padding="10")
        text_frame.grid(row=6, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        
        self.text_area = scrolledtext.ScrolledText(text_frame, width=100, height=20,
                                                  font=("Consolas", 10))
        self.text_area.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Status
        status_frame = ttk.Frame(main_frame)
        status_frame.grid(row=7, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        
        self.status_var = tk.StringVar(value="Pronto para começar")
        status_label = ttk.Label(status_frame, textvariable=self.status_var)
        status_label.grid(row=0, column=0, sticky=tk.W)
        
        # Configurar weights para responsividade
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(6, weight=1)
        file_frame.columnconfigure(0, weight=1)
        text_frame.columnconfigure(0, weight=1)
        text_frame.rowconfigure(0, weight=1)
        
        # Desabilitar botões inicialmente
        self.transcribe_btn.config(state='disabled')
        self.export_btn.config(state='disabled')
        
        # Configurar estilo para botão principal
        style = ttk.Style()
        style.configure("Accent.TButton", font=('Arial', 10, 'bold'))
    
    def select_audio_file(self):
        file_path = filedialog.askopenfilename(
            title="Selecionar Arquivo de Áudio",
            filetypes=[
                ("Arquivos de Áudio", "*.mp3 *.wav *.m4a *.flac *.aac *.ogg *.wma"),
                ("Todos os arquivos", "*.*")
            ]
        )
        
        if file_path:
            self.audio_path = file_path
            file_name = Path(file_path).name
            file_size = os.path.getsize(file_path) / (1024 * 1024)  # MB
            
            self.file_label.config(text=f"{file_name} ({file_size:.1f} MB)")
            self.transcribe_btn.config(state='normal')
            self.text_area.delete(1.0, tk.END)
            self.status_var.set(f"Arquivo selecionado: {file_name}")
    
    def start_transcription(self):
        if not self.audio_path:
            messagebox.showerror("Erro", "Por favor, selecione um arquivo de áudio.")
            return
        
        # Verificar se o arquivo é muito grande
        file_size = os.path.getsize(self.audio_path) / (1024 * 1024)
        if file_size > 100:  # 100MB
            if not messagebox.askyesno("Arquivo Grande", 
                                     f"O arquivo tem {file_size:.1f} MB. "
                                     "A transcrição pode demorar. Continuar?"):
                return
        
        self.progress.grid()
        self.progress['value'] = 0
        self.transcribe_btn.config(state='disabled')
        self.export_btn.config(state='disabled')
        self.status_var.set("Iniciando transcrição...")
        
        thread = threading.Thread(target=self.transcription_worker, 
                                args=(self.audio_path, self.model_var.get(), self.language_var.get()))
        thread.daemon = True
        thread.start()
    
    def transcription_worker(self, audio_path, model_size, language):
        try:
            self.update_status("Carregando modelo Whisper...")
            self.update_progress(10)
            
            # Importar whisper
            try:
                import whisper
            except ImportError:
                self.update_status("Instalando Whisper...")
                self.install_whisper()
                import whisper
            
            self.update_status("Carregando modelo de áudio...")
            self.update_progress(30)
            
            # Carregar modelo Whisper
            model = whisper.load_model(model_size)
            
            self.update_status("Processando áudio...")
            self.update_progress(50)
            
            # Configurar opções de transcrição
            options = {}
            if language:
                options["language"] = language
                options["task"] = "transcribe"
            
            # Transcrever áudio
            result = model.transcribe(audio_path, **options)
            
            self.update_status("Finalizando transcrição...")
            self.update_progress(90)
            
            transcription_text = result["text"]
            
            # Adicionar informações adicionais se disponíveis
            if "segments" in result:
                segments_text = "\n\n--- Segmentos de Áudio ---\n"
                for i, segment in enumerate(result["segments"]):
                    segments_text += f"\n[{i+1}] {segment['text']}\n"
                    if 'start' in segment and 'end' in segment:
                        segments_text += f"   Tempo: {segment['start']:.1f}s - {segment['end']:.1f}s\n"
                
                transcription_text += segments_text
            
            self.update_progress(100)
            self.root.after(0, lambda: self.on_transcription_finished(transcription_text))
            
        except Exception as e:
            self.root.after(0, lambda: self.on_transcription_error(str(e)))
    
    def install_whisper(self):
        """Instalar whisper se não estiver disponível"""
        try:
            import subprocess
            import sys
            
            self.update_status("Instalando OpenAI Whisper...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", "openai-whisper"])
            self.update_status("Whisper instalado com sucesso!")
            
        except Exception as e:
            raise Exception(f"Falha ao instalar Whisper: {str(e)}")
    
    def update_progress(self, value):
        self.root.after(0, lambda: self.progress.config(value=value))
    
    def update_status(self, message):
        self.root.after(0, lambda: self.status_var.set(message))
    
    def on_transcription_finished(self, text):
        self.transcription_text = text
        self.text_area.delete(1.0, tk.END)
        self.text_area.insert(1.0, text)
        self.export_btn.config(state='normal')
        self.transcribe_btn.config(state='normal')
        self.progress.grid_remove()
        self.status_var.set("✓ Transcrição concluída com sucesso!")
        
        # Mostrar estatísticas
        word_count = len(text.split())
        char_count = len(text)
        messagebox.showinfo("Sucesso", 
                          f"Transcrição concluída!\n\n"
                          f"Palavras: {word_count}\n"
                          f"Caracteres: {char_count}\n"
                          f"Modelo usado: {self.model_var.get()}")
    
    def on_transcription_error(self, error_message):
        self.transcribe_btn.config(state='normal')
        self.progress.grid_remove()
        self.status_var.set(f"✗ Erro: {error_message}")
        
        # Sugerir soluções para erros comuns
        if "whisper" in error_message.lower() or "model" in error_message.lower():
            error_message += "\n\nTente:\n1. Usar um modelo menor (tiny ou base)\n2. Verificar sua conexão com internet\n3. Reiniciar o programa"
        
        messagebox.showerror("Erro na Transcrição", error_message)
    
    def export_to_word(self):
        if not self.transcription_text:
            messagebox.showerror("Erro", "Nenhuma transcrição para exportar.")
            return
        
        file_path = filedialog.asksaveasfilename(
            title="Salvar Documento Word",
            defaultextension=".docx",
            filetypes=[("Documentos Word", "*.docx"), ("Todos os arquivos", "*.*")],
            initialfile=f"transcricao_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        )
        
        if file_path:
            try:
                self.status_var.set("Exportando para Word...")
                self.create_word_document(file_path, self.transcription_text)
                messagebox.showinfo("Sucesso", f"Documento salvo em:\n{file_path}")
                self.status_var.set("✓ Documento exportado com sucesso!")
                
                # Perguntar se quer abrir o arquivo
                if messagebox.askyesno("Abrir Documento", "Deseja abrir o documento agora?"):
                    os.startfile(file_path)  # Windows
                    
            except Exception as e:
                messagebox.showerror("Erro", f"Erro ao salvar documento:\n{str(e)}")
                self.status_var.set(f"✗ Erro na exportação: {str(e)}")
    
    def create_word_document(self, file_path, transcription_text):
        doc = Document()
        
        # Título
        title = doc.add_heading('TRANSCRIÇÃO DE ÁUDIO', 0)
        
        # Metadados em uma tabela
        doc.add_paragraph()  # Espaço
        
        metadata_table = doc.add_table(rows=4, cols=2)
        metadata_table.style = 'Light Grid Accent 1'
        
        # Preencher metadados
        metadata_table.cell(0, 0).text = "Data da Transcrição:"
        metadata_table.cell(0, 1).text = datetime.now().strftime('%d/%m/%Y às %H:%M')
        
        metadata_table.cell(1, 0).text = "Arquivo Original:"
        if self.audio_path:
            metadata_table.cell(1, 1).text = Path(self.audio_path).name
        
        metadata_table.cell(2, 0).text = "Modelo Whisper:"
        metadata_table.cell(2, 1).text = self.model_var.get()
        
        metadata_table.cell(3, 0).text = "Idioma:"
        metadata_table.cell(3, 1).text = self.language_var.get() if self.language_var.get() else "Auto-detectado"
        
        doc.add_paragraph()  # Espaço
        
        # Transcrição
        doc.add_heading('TRANSCRIÇÃO', level=1)
        
        # Adicionar linha divisória
        doc.add_paragraph("_" * 50)
        
        # Formatar o texto em parágrafos (CORREÇÃO AQUI)
        paragraphs = transcription_text.split('\n')
        for i, paragraph in enumerate(paragraphs):
            if paragraph.strip():
                # Destaque para títulos de seções
                if "---" in paragraph and "---" in paragraph:
                    p = doc.add_paragraph(paragraph.strip())
                    p.style = 'Heading 2'
                else:
                    p = doc.add_paragraph(paragraph.strip())
                
                # CORREÇÃO: Usar número inteiro diretamente em vez de tk.IntVar
                # Adicionar espaço entre parágrafos, mas não após o último
                if i < len(paragraphs) - 1:
                    p.paragraph_format.space_after = 6  # Número inteiro diretamente
        
        # Rodapé
        doc.add_page_break()
        footer = doc.add_paragraph()
        footer.add_run("Documento gerado automaticamente pelo ").bold = False
        footer.add_run("Conversor de Áudio para Word").bold = True
        footer.add_run(" usando OpenAI Whisper.").bold = False
        footer.alignment = 2  # Centralizado
        
        doc.save(file_path)
    
    def clear_all(self):
        self.audio_path = None
        self.transcription_text = ""
        self.file_label.config(text="Nenhum arquivo selecionado")
        self.text_area.delete(1.0, tk.END)
        self.transcribe_btn.config(state='disabled')
        self.export_btn.config(state='disabled')
        self.status_var.set("Pronto para começar")
        self.progress.grid_remove()
        self.progress['value'] = 0


def check_dependencies():
    """Verificar e instalar dependências necessárias"""
    try:
        import whisper
        print("✓ Whisper está instalado")
    except ImportError:
        print("Instalando Whisper...")
        import subprocess
        import sys
        subprocess.check_call([sys.executable, "-m", "pip", "install", "openai-whisper"])
    
    try:
        from docx import Document
        print("✓ python-docx está instalado")
    except ImportError:
        print("Instalando python-docx...")
        import subprocess
        import sys
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])


def main():
    # Verificar dependências
    print("Verificando dependências...")
    check_dependencies()
    
    # Criar interface
    root = tk.Tk()
    app = AudioToWordConverter(root)
    
    # Centralizar janela
    root.update_idletasks()
    x = (root.winfo_screenwidth() // 2) - (root.winfo_width() // 2)
    y = (root.winfo_screenheight() // 2) - (root.winfo_height() // 2)
    root.geometry(f"+{x}+{y}")
    
    root.mainloop()


if __name__ == "__main__":
    main()